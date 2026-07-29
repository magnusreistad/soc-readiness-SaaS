"""
app/utils/document_processor.py

Zero-storage document processor for the evidence evaluator.

Design principles:
    - Files are NEVER written to disk. All processing happens in memory.
    - Extracted text is passed directly to the AI pipeline and then
      discarded — it is never stored in the database.
    - File bytes are deleted from memory as soon as text is extracted.
    - Maximum file size is enforced before any processing begins.
    - Supported formats: PDF, images (PNG/JPG/WEBP), Excel (XLSX/XLS/CSV),
      Word (DOCX/DOC/RTF).

Word extraction strategy (layered fallbacks):
    1. python-docx  — handles well-formed .docx (OOXML/ZIP)
    2. RTF detection — handles files that are actually RTF
    3. Binary text   — extracts printable ASCII from any binary as last resort

Excel extraction strategy (layered fallbacks):
    1. openpyxl     — handles .xlsx (OOXML)
    2. xlrd         — handles old .xls binary format (if installed)
    3. CSV fallback  — tries to read as CSV if both fail

Dependencies:
    pip install pymupdf pillow python-docx openpyxl
    Optional: pip install xlrd    (for old .xls support)
"""

import csv
import io
import logging
import re
from dataclasses import dataclass
from typing import Literal

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Limits
# ---------------------------------------------------------------------------

MAX_FILE_SIZE_BYTES = 20 * 1024 * 1024   # 20 MB hard limit
MAX_EXTRACTED_CHARS = 50_000              # truncate extracted text at 50k chars
MAX_PAGES_PDF       = 50                  # only process first 50 pages of a PDF

ALLOWED_MIME_TYPES = {
    'application/pdf',
    'image/png',
    'image/jpeg',
    'image/webp',
    'image/gif',
    # Excel — both modern and legacy
    'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',  # .xlsx
    'application/vnd.ms-excel',                                            # .xls
    # CSV
    'text/csv',
    # Word — both modern and legacy
    'application/vnd.openxmlformats-officedocument.wordprocessingml.document',  # .docx
    'application/msword',                                                         # .doc
    'application/vnd.ms-word',                                                    # .doc (alt)
    'text/rtf',                                                                    # .rtf
    'application/rtf',                                                             # .rtf (alt)
    # Plain text
    'text/plain',
    # Catch-all for browsers that send octet-stream for known extensions
    'application/octet-stream',
}

FileType = Literal['pdf', 'image', 'excel', 'word', 'csv', 'text']


@dataclass
class ExtractionResult:
    text: str
    file_type: FileType
    page_count: int | None
    char_count: int
    was_truncated: bool
    error: str | None = None


# ---------------------------------------------------------------------------
# Public entry point
# ---------------------------------------------------------------------------

def extract_text(
    file_bytes: bytes,
    filename: str,
    mime_type: str,
) -> ExtractionResult:
    """
    Extracts text from an in-memory file without writing to disk.

    Args:
        file_bytes: Raw file content as bytes (from upload, never from disk).
        filename:   Original filename — used for extension fallback and logging.
        mime_type:  MIME type declared by the client. May be 'application/octet-stream'
                    if the browser doesn't know the type — we fall back to extension.

    Returns:
        ExtractionResult with extracted text and metadata.
        On failure, returns result with error set and empty text.
    """
    if len(file_bytes) > MAX_FILE_SIZE_BYTES:
        return ExtractionResult(
            text='', file_type='text', page_count=None, char_count=0,
            was_truncated=False,
            error=f'File too large. Maximum size is {MAX_FILE_SIZE_BYTES // (1024*1024)} MB.',
        )

    # Resolve effective MIME type — fall back to extension sniffing for octet-stream
    effective_mime = _resolve_mime(mime_type, filename, file_bytes)

    if effective_mime not in ALLOWED_MIME_TYPES:
        return ExtractionResult(
            text='', file_type='text', page_count=None, char_count=0,
            was_truncated=False,
            error=f'Unsupported file type: {mime_type}. '
                  f'Accepted: PDF, images, Excel (.xlsx/.xls), Word (.docx/.doc), CSV, plain text.',
        )

    try:
        if effective_mime == 'application/pdf':
            return _extract_pdf(file_bytes)

        elif effective_mime in ('image/png', 'image/jpeg', 'image/webp', 'image/gif'):
            return _extract_image(file_bytes, effective_mime)

        elif effective_mime in (
            'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            'application/vnd.ms-excel',
        ):
            return _extract_excel(file_bytes, filename)

        elif effective_mime == 'text/csv':
            return _extract_csv(file_bytes)

        elif effective_mime in (
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'application/msword',
            'application/vnd.ms-word',
            'text/rtf',
            'application/rtf',
        ):
            return _extract_word(file_bytes, filename)

        else:
            return _extract_plaintext(file_bytes)

    except Exception as exc:
        logger.error('Document extraction failed for %s: %s', filename, exc)
        return ExtractionResult(
            text='', file_type='text', page_count=None, char_count=0,
            was_truncated=False,
            error=f'Could not extract text from this file: {exc}',
        )
    finally:
        del file_bytes


# ---------------------------------------------------------------------------
# MIME resolution — handle octet-stream by sniffing extension and magic bytes
# ---------------------------------------------------------------------------

_EXT_MIME = {
    '.pdf':  'application/pdf',
    '.png':  'image/png',
    '.jpg':  'image/jpeg',
    '.jpeg': 'image/jpeg',
    '.webp': 'image/webp',
    '.gif':  'image/gif',
    '.xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
    '.xls':  'application/vnd.ms-excel',
    '.csv':  'text/csv',
    '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    '.doc':  'application/msword',
    '.rtf':  'text/rtf',
    '.txt':  'text/plain',
}


def _resolve_mime(mime_type: str, filename: str, file_bytes: bytes) -> str:
    """
    Returns the best MIME type to use for extraction.
    If the browser sends 'application/octet-stream', we try the filename
    extension first, then magic bytes.
    """
    if mime_type and mime_type != 'application/octet-stream':
        return mime_type

    # Try filename extension
    lower = filename.lower()
    for ext, resolved in _EXT_MIME.items():
        if lower.endswith(ext):
            logger.info(
                'Resolved octet-stream to %s via extension for %s', resolved, filename
            )
            return resolved

    # Try magic bytes
    if file_bytes[:4] == b'%PDF':
        return 'application/pdf'
    if file_bytes[:4] in (b'PK\x03\x04',):
        # ZIP-based: could be .docx or .xlsx — check internal structure
        if b'word/' in file_bytes[:2000]:
            return 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        if b'xl/' in file_bytes[:2000]:
            return 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    if file_bytes[:5] == b'{\\rtf':
        return 'text/rtf'

    return mime_type  # give up — let the MIME check catch it


# ---------------------------------------------------------------------------
# Format-specific extractors
# ---------------------------------------------------------------------------

def _truncate(text: str) -> tuple[str, bool]:
    if len(text) <= MAX_EXTRACTED_CHARS:
        return text, False
    return text[:MAX_EXTRACTED_CHARS], True


def _extract_pdf(file_bytes: bytes) -> ExtractionResult:
    try:
        import fitz
    except ImportError:
        return ExtractionResult(
            text='', file_type='pdf', page_count=None, char_count=0,
            was_truncated=False,
            error='PDF processing requires PyMuPDF: pip install pymupdf',
        )

    doc = fitz.open(stream=file_bytes, filetype='pdf')
    pages_to_read = min(len(doc), MAX_PAGES_PDF)
    parts = []
    for i in range(pages_to_read):
        page = doc[i]
        parts.append(page.get_text())
        page = None
    doc.close()

    raw = '\n'.join(parts)
    text, truncated = _truncate(raw)
    return ExtractionResult(
        text=text, file_type='pdf', page_count=pages_to_read,
        char_count=len(text), was_truncated=truncated,
    )


def _extract_image(file_bytes: bytes, mime_type: str) -> ExtractionResult:
    """
    Images are passed as base64 to the vision AI — no OCR needed here.
    Returns the base64 string as text so the AI router can handle it.
    """
    import base64
    b64 = base64.b64encode(file_bytes).decode('utf-8')
    return ExtractionResult(
        text=b64, file_type='image', page_count=1,
        char_count=len(b64), was_truncated=False,
    )


def _extract_excel(file_bytes: bytes, filename: str) -> ExtractionResult:
    """
    Layered Excel extraction:
      1. openpyxl  — .xlsx (OOXML)
      2. xlrd      — old .xls binary (if installed)
      3. CSV       — last resort if the bytes look like text
    """
    # ── Strategy 1: openpyxl for .xlsx ──────────────────────────────────────
    try:
        import openpyxl
        wb = openpyxl.load_workbook(
            io.BytesIO(file_bytes), read_only=True, data_only=True,
        )
        parts = []
        sheet_count = 0
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            sheet_count += 1
            parts.append(f'[Sheet: {sheet_name}]')
            for row in ws.iter_rows(values_only=True):
                cells = [str(c) if c is not None else '' for c in row]
                if any(c.strip() for c in cells):
                    parts.append('\t'.join(cells))
        wb.close()

        raw = '\n'.join(parts)
        if raw.strip():
            text, truncated = _truncate(raw)
            return ExtractionResult(
                text=text, file_type='excel', page_count=sheet_count,
                char_count=len(text), was_truncated=truncated,
            )
    except Exception as e:
        logger.info('openpyxl failed for %s: %s — trying xlrd', filename, e)

    # ── Strategy 2: xlrd for old .xls ────────────────────────────────────────
    try:
        import xlrd  # type: ignore
        wb = xlrd.open_workbook(file_contents=file_bytes)
        parts = []
        for sheet in wb.sheets():
            parts.append(f'[Sheet: {sheet.name}]')
            for row_idx in range(sheet.nrows):
                cells = [str(sheet.cell_value(row_idx, col)) for col in range(sheet.ncols)]
                if any(c.strip() for c in cells):
                    parts.append('\t'.join(cells))
        raw = '\n'.join(parts)
        if raw.strip():
            text, truncated = _truncate(raw)
            return ExtractionResult(
                text=text, file_type='excel', page_count=wb.nsheets,
                char_count=len(text), was_truncated=truncated,
            )
    except ImportError:
        logger.debug('xlrd not installed — skipping .xls fallback')
    except Exception as e:
        logger.info('xlrd failed for %s: %s — trying CSV fallback', filename, e)

    # ── Strategy 3: CSV fallback ──────────────────────────────────────────────
    try:
        return _extract_csv(file_bytes)
    except Exception:
        pass

    return ExtractionResult(
        text='', file_type='excel', page_count=None, char_count=0,
        was_truncated=False,
        error=(
            'Could not extract text from this spreadsheet. '
            'If this is an old .xls file, try saving as .xlsx first. '
            'Alternatively, export as CSV.'
        ),
    )


def _extract_csv(file_bytes: bytes) -> ExtractionResult:
    text_io = io.StringIO(file_bytes.decode('utf-8', errors='replace'))
    reader = csv.reader(text_io)
    rows = ['\t'.join(row) for row in reader if any(cell.strip() for cell in row)]
    raw = '\n'.join(rows)
    text, truncated = _truncate(raw)
    return ExtractionResult(
        text=text, file_type='csv', page_count=None,
        char_count=len(text), was_truncated=truncated,
    )


def _extract_word(file_bytes: bytes, filename: str) -> ExtractionResult:
    """
    Layered Word extraction:
      1. python-docx   — well-formed .docx (OOXML/ZIP with correct metadata)
      2. Direct ZIP/XML — .docx missing [Content_Types].xml or other metadata
                          (Pages exports, LibreOffice quirks, repacked files)
      3. RTF parser    — files that are actually RTF
      4. Binary text   — printable ASCII extraction from any binary as last resort
    """
    # ── Strategy 1: python-docx (standard .docx) ────────────────────────────
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        parts = [para.text for para in doc.paragraphs if para.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append('\t'.join(cells))
        raw = '\n'.join(parts)
        if raw.strip():
            text, truncated = _truncate(raw)
            return ExtractionResult(
                text=text, file_type='word', page_count=None,
                char_count=len(text), was_truncated=truncated,
            )
    except ImportError:
        return ExtractionResult(
            text='', file_type='word', page_count=None, char_count=0,
            was_truncated=False,
            error='Word processing requires python-docx: pip install python-docx',
        )
    except Exception as e:
        logger.info('python-docx failed for %s: %s — trying direct ZIP extraction', filename, e)

    # ── Strategy 2: Direct ZIP + XML extraction ──────────────────────────────
    # Handles the exact '[Content_Types].xml' error — python-docx requires this
    # metadata manifest, but the actual text lives in word/document.xml and we
    # can read it directly with zipfile without caring about the manifest.
    # Triggered by: Pages exports, LibreOffice-created files, repacked .docx,
    # or any OOXML ZIP that's missing its content type declarations.
    if file_bytes[:2] == b'PK':    # ZIP magic bytes
        try:
            import zipfile
            from xml.etree import ElementTree as ET

            with zipfile.ZipFile(io.BytesIO(file_bytes)) as zf:
                names = zf.namelist()

                # Find word/document.xml (handles both standard and renamed paths)
                doc_xml_name = next(
                    (n for n in names if n.endswith('document.xml') and 'word' in n),
                    None,
                )
                if not doc_xml_name:
                    # Some files use 'Document.xml' (capital D)
                    doc_xml_name = next(
                        (n for n in names if n.lower().endswith('document.xml')),
                        None,
                    )

                if doc_xml_name:
                    xml_bytes = zf.read(doc_xml_name)
                    root = ET.fromstring(xml_bytes)

                    # Extract all <w:t> text elements — this is the actual content
                    # Namespace used by OOXML for Word content
                    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                    texts = [el.text for el in root.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t') if el.text]

                    if texts:
                        raw = ' '.join(t for t in texts if t.strip())
                        if raw.strip():
                            text, truncated = _truncate(raw)
                            return ExtractionResult(
                                text=text, file_type='word', page_count=None,
                                char_count=len(text), was_truncated=truncated,
                            )

        except zipfile.BadZipFile:
            pass    # not a ZIP at all — fall through to RTF/binary
        except Exception as e:
            logger.info('Direct ZIP/XML extraction failed for %s: %s', filename, e)

    # ── Strategy 3: RTF extraction ────────────────────────────────────────────
    if file_bytes[:4] == b'{\\rt' or file_bytes[:4] == b'{\x5crt':
        try:
            raw_str = file_bytes.decode('latin-1', errors='replace')
            text_str = re.sub(r'\\[a-zA-Z]+[-\d]* ?', ' ', raw_str)
            text_str = re.sub(r'[{}\\]', ' ', text_str)
            text_str = re.sub(r'\s+', ' ', text_str).strip()
            if len(text_str) > 50:
                text, truncated = _truncate(text_str)
                return ExtractionResult(
                    text=text, file_type='word', page_count=None,
                    char_count=len(text), was_truncated=truncated,
                )
        except Exception as e:
            logger.info('RTF extraction failed for %s: %s', filename, e)

    # ── Strategy 4: Binary text extraction ───────────────────────────────────
    try:
        raw_str = file_bytes.decode('latin-1', errors='replace')
        chunks = re.findall(r'[A-Za-z0-9 \t.,;:!?()\-\'\"@/]{6,}', raw_str)
        readable = [
            c.strip() for c in chunks
            if sum(1 for ch in c if ch.isalpha()) > len(c) * 0.5
        ]
        combined = ' '.join(readable)
        if len(combined) > 80:
            text, truncated = _truncate(combined)
            note = (
                '[Note: This document could not be parsed as a standard .docx file. '
                'Text was extracted from the binary — some content may be missing. '
                'For best results, save as .docx from Microsoft Word or LibreOffice.]\n\n'
            )
            return ExtractionResult(
                text=note + text, file_type='word', page_count=None,
                char_count=len(text), was_truncated=truncated,
            )
    except Exception as e:
        logger.info('Binary text extraction failed for %s: %s', filename, e)

    return ExtractionResult(
        text='', file_type='word', page_count=None, char_count=0,
        was_truncated=False,
        error=(
            'Could not extract text from this Word document. '
            'The file may be in an unsupported format. '
            'Try: File → Save As → .docx in Microsoft Word or LibreOffice, '
            'or export as PDF.'
        ),
    )


def _extract_plaintext(file_bytes: bytes) -> ExtractionResult:
    raw = file_bytes.decode('utf-8', errors='replace')
    text, truncated = _truncate(raw)
    return ExtractionResult(
        text=text, file_type='text', page_count=None,
        char_count=len(text), was_truncated=truncated,
    )